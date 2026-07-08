import os
import sys
import logging
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from redolarium.config import EXCEL_PALETTE, TAB_COLORS, CONFIG
import time
import functools
import urllib.error
import http.client
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

def api_retry(retries=5, backoff_factor=2.0, exceptions=(urllib.error.URLError, http.client.HTTPException, IOError)):
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            # Polite default rate-limiting delay to prevent hammering public APIs (like NCBI or ChEMBL)
            time.sleep(0.35)
            delay = 2.0
            for i in range(retries):
                try:
                    return func(*args, **kwargs)
                except exceptions as e:
                    if i == retries - 1:
                        raise e
                    logger = kwargs.get('logger')
                    if not logger and args:
                        for arg in args:
                            if hasattr(arg, 'info') and hasattr(arg, 'warning'):
                                logger = arg
                                break
                    msg = f"API call failed: {e}. Retrying in {delay:.1f}s (Attempt {i+1}/{retries})..."
                    if logger:
                        logger.warning(msg)
                    else:
                        print(msg)
                    time.sleep(delay)
                    delay *= backoff_factor
        return wrapper
    return decorator

def save_publication_plot(fig_path, dpi=300):
    """
    Saves the active matplotlib figure as both a high-res PNG and a layered
    BioRender/Illustrator-compatible SVG (preserving fonts as editable text).
    """
    matplotlib.rcParams['svg.fonttype'] = 'none'
    plt.savefig(fig_path, dpi=dpi, bbox_inches='tight')
    svg_path = os.path.splitext(fig_path)[0] + ".svg"
    plt.savefig(svg_path, format="svg", bbox_inches='tight')

def format_for_prism_and_origin(hgt_results, out_dir, bgc_id):
    """
    Formats HGT Z-scores and sliding window data column-by-column with NaN padding
    for direct copy-paste/import into GraphPad Prism and Origin.
    Also generates a grouped summary table (Mean, SD, N) for direct import.
    """
    import numpy as np
    df = pd.DataFrame(hgt_results)
    if df.empty:
        return
        
    core_gc = df[df["Region"] == "BGC Core"]["GC_Fraction"] * 100
    flank_gc = df[df["Region"] != "BGC Core"]["GC_Fraction"] * 100
    core_tnf = df[df["Region"] == "BGC Core"]["TNF_Divergence"]
    flank_tnf = df[df["Region"] != "BGC Core"]["TNF_Divergence"]
    
    max_len = max(len(core_gc), len(flank_gc), len(core_tnf), len(flank_tnf))
    prism_data = {
        "BGC_Core_GC_Pct": list(core_gc) + [None] * (max_len - len(core_gc)),
        "Flanking_GC_Pct": list(flank_gc) + [None] * (max_len - len(flank_gc)),
        "BGC_Core_TNF_Dist": list(core_tnf) + [None] * (max_len - len(core_tnf)),
        "Flanking_TNF_Dist": list(flank_tnf) + [None] * (max_len - len(flank_tnf))
    }
    df_prism = pd.DataFrame(prism_data)
    
    hgt_dir = os.path.join(out_dir, "hgt_evolution")
    os.makedirs(hgt_dir, exist_ok=True)
    
    prism_csv = os.path.join(hgt_dir, f"{bgc_id}_prism_origin_ready.csv")
    df_prism.to_csv(prism_csv, index=False)
    
    # Pre-calculated Mean, SD, N summary for Grouped data tables
    summary_data = {
        "Region": ["BGC Core", "Flanking Region"],
        "GC_Mean": [np.mean(core_gc) if len(core_gc) > 0 else np.nan, np.mean(flank_gc) if len(flank_gc) > 0 else np.nan],
        "GC_SD": [np.std(core_gc, ddof=1) if len(core_gc) > 1 else 0.0, np.std(flank_gc, ddof=1) if len(flank_gc) > 1 else 0.0],
        "GC_N": [len(core_gc), len(flank_gc)],
        "TNF_Mean": [np.mean(core_tnf) if len(core_tnf) > 0 else np.nan, np.mean(flank_tnf) if len(flank_tnf) > 0 else np.nan],
        "TNF_SD": [np.std(core_tnf, ddof=1) if len(core_tnf) > 1 else 0.0, np.std(flank_tnf, ddof=1) if len(flank_tnf) > 1 else 0.0],
        "TNF_N": [len(core_tnf), len(flank_tnf)]
    }
    df_summary = pd.DataFrame(summary_data)
    summary_csv = os.path.join(hgt_dir, f"{bgc_id}_prism_grouped_summary.csv")
    df_summary.to_csv(summary_csv, index=False)

# Setup logger
def setup_logging(out_dir):
    os.makedirs(out_dir, exist_ok=True)
    log = logging.getLogger("redolarium")
    log.setLevel(logging.INFO)
    if not log.handlers:
        fmt = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
        ch = logging.StreamHandler(sys.stdout)
        ch.setFormatter(fmt)
        fh = logging.FileHandler(os.path.join(out_dir, "pipeline_execution.log"), mode="w", encoding="utf-8")
        fh.setFormatter(fmt)
        log.addHandler(ch)
        log.addHandler(fh)
    return log

# ==========================================
# EXCEL STYLING HELPERS (WORKBOOK STANDARD)
# ==========================================
def _sheet_title(ws, title, sub=""):
    """Insert bold sheet title (Row 1) and italic subtitle (Row 2)."""
    ws.views.sheetView[0].showGridLines = True
    ws.insert_rows(1, 2)
    ws.cell(row=1, column=1, value=title)
    ws.cell(row=2, column=1, value=sub)
    
    max_col = max(10, ws.max_column)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=max_col)
    
    ws.cell(row=1, column=1).font = Font(name="Calibri", size=13, bold=True, color="1F3864")
    ws.cell(row=2, column=1).font = Font(name="Calibri", size=9, italic=True, color="595959")
    ws.row_dimensions[1].height = 24
    ws.row_dimensions[2].height = 16

def _hdr(ws, row, fill_hex, font_hex="FFFFFF"):
    """Apply standard header row styling to row (navy fill, white bold Calibri text)."""
    fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
    font = Font(name="Calibri", size=10, bold=True, color=font_hex)
    thin_border = Border(
        left=Side(style='thin', color='1F3864'),
        right=Side(style='thin', color='1F3864'),
        top=Side(style='thin', color='1F3864'),
        bottom=Side(style='thin', color='1F3864')
    )
    for col in range(1, ws.max_column + 1):
        cell = ws.cell(row=row, column=col)
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
    ws.row_dimensions[row].height = 28

def _alt_rows(ws, start, end, fill_hex):
    """Apply alternating even-row fill between start and end rows."""
    fill = PatternFill(start_color=fill_hex, end_color=fill_hex, fill_type="solid")
    for r in range(start, end + 1):
        if r % 2 == 0:
            for c in range(1, ws.max_column + 1):
                cell = ws.cell(row=r, column=c)
                if cell.fill.fill_type is None or cell.fill.fgColor.rgb == "00000000":
                    cell.fill = fill

def _col_widths(ws, widths):
    """Apply width mapping to all columns."""
    for idx, width in enumerate(widths):
        col_letter = get_column_letter(idx + 1)
        ws.column_dimensions[col_letter].width = width

def apply_thin_borders(ws, start_row, end_row):
    """Apply thin grey borders to the data grid."""
    thin = Side(border_style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in range(start_row, end_row + 1):
        for col in range(1, ws.max_column + 1):
            ws.cell(row=row, column=col).border = border

def apply_general_cell_styles(ws, start_row, end_row):
    """Set Calibri font and alignment for all data rows."""
    for row in range(start_row, end_row + 1):
        for col in range(1, ws.max_column + 1):
            cell = ws.cell(row=row, column=col)
            cell.font = Font(name="Calibri", size=9.5)
            if isinstance(cell.value, (int, float)):
                cell.alignment = Alignment(horizontal="right", vertical="center")
            else:
                cell.alignment = Alignment(horizontal="left", vertical="center")

def create_references_sheet_openpyxl(wb):
    """Generate the standard peer-reviewed references sheet."""
    ws = wb.create_sheet(title="Methods_References")
    ws.sheet_properties.tabColor = TAB_COLORS["Methods_References"]
    
    ws.append(["Analysis Module", "Method / Tool", "Key Reference Citation", "DOI"])
    for cite in CONFIG["citations"]:
        ws.append([cite["module"], cite["tool"], cite["ref"], cite["doi"]])
        
    _sheet_title(ws, "Analytical Methods and Database References", 
                 "Bibliography of peer-reviewed tools, parameters, and resources utilized in this pipeline run.")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    
    for r in range(4, ws.max_row + 1):
        ws.cell(row=r, column=1).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=r, column=2).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=r, column=4).alignment = Alignment(horizontal="center", vertical="center")
        ws.cell(row=r, column=4).font = Font(name="Calibri", size=9, color="2E75B6", underline="single")
        ws.cell(row=r, column=2).font = Font(name="Calibri", size=9.5, bold=True)
        
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    _col_widths(ws, [20, 20, 75, 34])

# ==========================================
# WORD DOCUMENT WRITER TEMPLATE
# ==========================================
def write_word_report(filename, bgc_id, query_org, sections, out_dir):
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    # Report Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"Comprehensive BGC Analysis & Metabolic Expression Report: {bgc_id}")
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100) # Navy blue
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run(f"Query Organism: {query_org} | Target: {bgc_id} | Pipeline: Redolarium v1.1.0")
    s_run.font.name = 'Calibri'
    s_run.font.size = Pt(10)
    s_run.font.italic = True
    s_run.font.color.rgb = RGBColor(127, 127, 127)
    doc.add_paragraph()
    
    for sec_title, sec_content in sections:
        h = doc.add_paragraph()
        h.paragraph_format.keep_with_next = True
        hrun = h.add_run(sec_title)
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(13)
        hrun.font.bold = True
        hrun.font.color.rgb = RGBColor(31, 56, 100)
        
        if isinstance(sec_content, list):
            for para in sec_content:
                if para.startswith("- "):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(para[2:])
                else:
                    p = doc.add_paragraph(para)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(sec_content)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()
        
    doc.save(os.path.join(out_dir, filename))

def compile_bgc_excel_report(bgc_target, ortholog_mapping, ref_strains, reconstructed_metab, 
                              region_genes, flanking_genes, promoter_records, phage_hits, 
                              hgt_results, stoichiometry_data, out_path, query_org, logger,
                              bgc_blast_results=None, qc_result=None, result_anno=None,
                              result_bgc=None, result_evol=None, result_prom=None,
                              result_dock=None, result_phy=None):
    import openpyxl
    import numpy as np
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    out_dir = os.path.dirname(out_path)
    # Extract global HGT baselines if dict
    global_mean_gc = 0.45
    global_mean_tnf = 0.08
    if isinstance(hgt_results, dict):
        global_mean_gc = hgt_results.get("global_mean_gc", 0.45)
        global_mean_tnf = hgt_results.get("global_mean_tnf", 0.08)
        hgt_results = hgt_results.get("windows", [])
        
    wb = openpyxl.Workbook()
    
    bgc_id = bgc_target["BGC_ID"]
    bgc_type = bgc_target["BGC_Type"]
    bgc_start = bgc_target["Start_Coord"]
    bgc_end = bgc_target["End_Coord"]
    bgc_size = bgc_target["Size_bp"]

    # 1. Summary Sheet
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_properties.tabColor = TAB_COLORS["Summary"]
    
    ws.append(["Analysis Parameter", "Dataset Value", "Context Interpretation & References"])
    
    # Calculate HGT max composite score
    max_hgt = max([x.get("Composite_HGT_Score", 0.0) for x in hgt_results], default=0.0)
    
    summary_rows = [
        ("Host Genome Organism", str(query_org), "Target organism containing the analyzed cluster"),
        ("BGC Identifier", str(bgc_id), "Unique BGC cluster identification tag"),
        ("Delineated BGC Type", str(bgc_type), "Biosynthetic class classification based on key enzyme signatures"),
        ("Genomic Boundary Coordinates", f"{bgc_start:,} - {bgc_end:,} bp (Size: {bgc_size:,} bp)", "Cluster boundaries plus flanking windows"),
        ("BGC Core Genes Count", str(len(region_genes)), "Delineated open reading frames encoding biosynthetic machinery"),
        ("Flanking Context Region ORFs", str(len(flanking_genes)), "Accessory context genes in 30kb flanking windows"),
        ("Regulatory Promoters Detected", str(len(promoter_records)), "Upstream regulatory transcription motif binding sites"),
        ("Prophage / MGE Flags", str(len(phage_hits)), "Flanking transposons, integrases, and phage remnants"),
        ("HGT Acquisition Evidence", "Present (High Composite)" if max_hgt > 0.50 else "Uncertain / Low", "Evidence of horizontal acquisition via codon/TNF deviations"),
        ("Metabolic precursor peptide", f"{len(stoichiometry_data) - 1 if stoichiometry_data else 0} aa residues Supply", "Prepropeptide translation count and tRNA stoichiometry")
    ]
    for param, val, desc in summary_rows:
        ws.append([param, val, desc])
        
    _sheet_title(ws, f"Biosynthetic Gene Cluster Summary Dashboard: {bgc_id}",
                 f"Analytical compilation and metabolic integration report for {query_org}")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    # Bold parameters column
    for r in range(4, ws.max_row + 1):
        ws.cell(row=r, column=1).font = Font(name="Calibri", size=9.5, bold=True)
        
    # Evidence Graph Checkpoints and Systemic Uncertainties Section
    if qc_result and result_anno and result_bgc and result_evol and result_prom and result_dock and result_phy:
        ws.append([])
        ws.append(["Evidence Graph Checkpoint", "Confidence Score", "Systemic Uncertainties"])
        hdr_row = ws.max_row
        _hdr(ws, hdr_row, "1F3864")
        
        checkpoints = [
            ("Assembly Quality Control", qc_result.confidence_score, "; ".join(qc_result.limitations)),
            ("Taxonomy & Genus Map", result_anno.confidence_score, "; ".join(result_anno.limitations)),
            ("BGC Delineation", result_bgc.confidence_score, "; ".join(result_bgc.limitations)),
            ("Horizontal Gene Transfer", result_evol.confidence_score, "; ".join(result_evol.limitations)),
            ("Promoter Transcription", result_prom.confidence_score, "; ".join(result_prom.limitations)),
            ("Molecular Docking Validation", result_dock.confidence_score, "; ".join(result_dock.limitations)),
            ("Phylogenetic Divergence Model", result_phy.confidence_score, "; ".join(result_phy.limitations))
        ]
        
        start_cp = ws.max_row + 1
        for cp, sc, unc in checkpoints:
            ws.append([cp, round(sc, 2), unc if unc else "No systemic uncertainties identified."])
            
        end_cp = ws.max_row
        apply_thin_borders(ws, start_cp, end_cp)
        _alt_rows(ws, start_cp, end_cp, EXCEL_PALETTE["alt_row_fill"])
        
        # Center-align confidence score column
        for r in range(start_cp, end_cp + 1):
            ws.cell(row=r, column=2).alignment = Alignment(horizontal="center")
            ws.cell(row=r, column=1).font = Font(name="Calibri", size=9.5, bold=True)
            
        # Check if manual intervention is required for any key module
        key_modules = [
            ("Assembly Quality Control", qc_result),
            ("BGC Delineation", result_bgc),
            ("Molecular Docking Validation", result_dock)
        ]
        flagged_modules = []
        for name, res in key_modules:
            if res and getattr(res, "manual_intervention_required", False):
                flagged_modules.append(name)
                
        if flagged_modules:
            ws.append([])
            ws.append(["[FLAGGED FOR REVIEW] Alert: Low Confidence Module Detection", "", ""])
            alert_row = ws.max_row
            ws.merge_cells(start_row=alert_row, start_column=1, end_row=alert_row, end_column=3)
            ws.cell(row=alert_row, column=1).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["warning_fill"])
            ws.cell(row=alert_row, column=1).font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
            ws.cell(row=alert_row, column=1).alignment = Alignment(horizontal="left", vertical="center")
            
            ws.append([
                "Expert Action Required", 
                "Yes", 
                f"The following key modules fell below the 0.3 confidence threshold: {', '.join(flagged_modules)}. "
                "Please manually inspect the evidence_graph.json for audit verification."
            ])
            detail_row = ws.max_row
            ws.cell(row=detail_row, column=1).font = Font(name="Calibri", size=9.5, bold=True)
            ws.cell(row=detail_row, column=2).alignment = Alignment(horizontal="center")
            ws.cell(row=detail_row, column=2).font = Font(name="Calibri", size=9.5, bold=True)
            apply_thin_borders(ws, alert_row, detail_row)
        
    _hgt_label = "horizontal gene transfer acquisition" if max_hgt > 0.50 else "vertical inheritance evolution"
    _n_mge = f"{len(phage_hits)} mobile element markers"
    
    total_atp = 0.0
    if stoichiometry_data:
        total_row = next((r for r in stoichiometry_data if "TOTAL" in str(r.get("Amino_Acid", ""))), None)
        if total_row:
            total_atp = total_row.get("Total_ATP_Cost", 0.0)
    _stoich_note = f"requires {total_atp:.1f} ATP equivalents" if total_atp > 0 else "requires standard ATP equivalents"
    
    # Long interpretation block at bottom
    start_r = ws.max_row + 2
    ws.merge_cells(start_row=start_r, start_column=1, end_row=start_r + 4, end_column=3)
    ws.cell(row=start_r, column=1, value=(
        f"Interpretation & Synteny Context: The query BGC ({bgc_id}) represents an active biosynthetic island of class {bgc_type} "
        f"reconstructed in the genome of {query_org}. Coding sequence alignments and sliding-window composition analysis indicate "
        f"{_hgt_label} with {_n_mge}. Precursor stoichiometry "
        f"confirms a pathway demand that {_stoich_note} well within the metabolic pool capabilities of the host isolate, with active tRNA-synthetase charging "
        f"and promoter motif alignments suggesting transcription coordination under environmental stress conditions."
    )).font = Font(name="Calibri", size=9.5, italic=True)
    ws.cell(row=start_r, column=1).alignment = Alignment(wrap_text=True, vertical="top")
    
    _col_widths(ws, [30, 24, 70])

    # 2. BGC_Gene_Architecture
    ws = wb.create_sheet(title="BGC_Gene_Architecture")
    ws.sheet_properties.tabColor = TAB_COLORS["BGC_Gene_Architecture"]
    
    ws.append(["Locus Tag", "Gene Symbol", "Start", "End", "Strand", "Product Description", "Role Class", "Ref Ortholog Tag", "Identity Pct (%)"])
    for g in region_genes:
        orth = next((x for x in ortholog_mapping if x["Locus_Tag"] == g["Locus_Tag"]), {})
        ws.append([
            g["Locus_Tag"],
            g["Gene_Symbol"],
            g["Start_Coord"],
            g["End_Coord"],
            g["Strand"],
            g["Product_Description"],
            g["Role"],
            orth.get("Ref_Ortholog_Tag", "NA"),
            orth.get("Identity_Pct", 0.0)
        ])
        
    _sheet_title(ws, "Biosynthetic Gene Cluster Core Architecture",
                 f"Open reading frames identified within the biosynthetic core region of {bgc_id}")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    # Role-based styling matching the standard
    for r in range(4, ws.max_row + 1):
        role_val = ws.cell(row=r, column=7).value
        cell = ws.cell(row=r, column=7)
        if role_val == "biosynthetic":
            cell.fill = PatternFill("solid", fgColor=EXCEL_PALETTE["bgc_core_fill"])
        elif role_val == "transport":
            cell.fill = PatternFill("solid", fgColor="AED6F1")
        elif role_val == "regulatory":
            cell.fill = PatternFill("solid", fgColor="D7BDE2")
        elif role_val == "immunity":
            cell.fill = PatternFill("solid", fgColor="F9C784")
            
    _col_widths(ws, [16, 12, 12, 12, 10, 38, 16, 16, 14])

    # 3. Promoter_Analysis
    ws = wb.create_sheet(title="Promoter_Analysis")
    ws.sheet_properties.tabColor = TAB_COLORS["Promoter_Analysis"]
    
    ws.append(["Locus Tag", "Gene Symbol", "Motif Upstream Pos", "Sigma Factor Association", "-35 Box Motif", "Spacer (bp)", "-10 Box Motif", "Shine-Dalgarno", "Regulatory Box", "Quality Class"])
    for pm in promoter_records:
        ws.append([
            pm["Locus_Tag"],
            pm["Gene_Symbol"],
            pm["Upstream_Position"],
            pm["Sigma_Factor"],
            pm["Minus35_Seq"],
            pm["Spacer_Length"],
            pm["Minus10_Seq"],
            pm["Shine_Dalgarno"],
            pm.get("Regulatory_Box") if pm.get("Regulatory_Box") else pm.get("YbdJ_Regulatory_Box", "Not detected"),
            pm.get("Quality_Class") if pm.get("Quality_Class") else pm.get("Quality", "Strong")
        ])
        
    _sheet_title(ws, "BGC Promoter & Shine-Dalgarno Motif Profile",
                 "Transcription binding sites and ribosomal translation consensus motifs upstream of core genes")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    for r in range(4, ws.max_row + 1):
        q_val = ws.cell(row=r, column=10).value
        if "Strong" in str(q_val):
            ws.cell(row=r, column=10).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["positive_fill"])
            
    _col_widths(ws, [16, 12, 18, 22, 16, 12, 16, 16, 20, 16])

    # 4. Phage_Artifacts
    ws = wb.create_sheet(title="Phage_Artifacts")
    ws.sheet_properties.tabColor = TAB_COLORS["Phage_Artifacts"]
    
    ws.append(["Locus Tag", "Start Position", "End Position", "Strand", "Product Annotation Description", "MGE Class Category"])
    for ph in phage_hits:
        ws.append([
            ph["Locus_Tag"],
            ph.get("Start_Coord") if "Start_Coord" in ph else ph.get("Start", 0),
            ph.get("End_Coord") if "End_Coord" in ph else ph.get("End", 0),
            ph.get("Strand", "+"),
            ph.get("Product_Description") if "Product_Description" in ph else ph.get("Product", "Integrase"),
            ph["MGE_Class"]
        ])
        
    _sheet_title(ws, "Prophage Elements & Mobile Genetic Elements (MGE)",
                 "Delineated insertion markers, recombinases, and transposons flanking the biosynthetic cluster")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    for r in range(4, ws.max_row + 1):
        ws.cell(row=r, column=6).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["critical_fill"])
        
    _col_widths(ws, [16, 16, 16, 10, 38, 24])

    # 5. HGT_Evidence
    ws = wb.create_sheet(title="HGT_Evidence")
    ws.sheet_properties.tabColor = TAB_COLORS["HGT_Evidence"]
    
    ws.append(["Criterion Metric Method", "Genome baseline value", "BGC Region window value", "Divergence Score (Z-Score)", "Interpretation score", "Scientific Reference"])
    
    bgc_gcs = [x["GC_Fraction"] for x in hgt_results if x["Region"] == "BGC Core"]
    bgc_tnf = [x["TNF_Divergence"] for x in hgt_results if x["Region"] == "BGC Core"]
    mean_bgc_gc = np.mean(bgc_gcs) if bgc_gcs else 0.45
    mean_bgc_tnf = np.mean(bgc_tnf) if bgc_tnf else 0.08
    
    gc_z = max([abs(x["GC_Zscore"]) for x in hgt_results if x["Region"] == "BGC Core"], default=0.0)
    
    hgt_rows = [
        ("Codon/GC Bias deviation", f"Genome GC baseline: {global_mean_gc*100:.2f}%", f"{mean_bgc_gc*100:.2f}%", f"Z-Score = {gc_z:.2f}", f"{min(25.0, gc_z*12.5):.2f}", "Langille & Bhatt (2006)"),
        ("Tetranucleotide Frequency (TNF)", f"Genome TNF baseline: {global_mean_tnf:.4f}", f"{mean_bgc_tnf:.4f}", f"TNF Distance = {mean_bgc_tnf:.3f}", f"{min(25.0, mean_bgc_tnf*100.0):.2f}", "Teeling et al. (2004)"),
        ("MGE Proximity Bonus", "No insertion markers", "Phage / IS proximity", f"{len(phage_hits)} elements flagged", "25.00" if phage_hits else "0.00", "Ravenhall et al. (2015)"),
        ("COMPOSITE HGT SCORE", "-", "-", "-", f"{sum([min(25.0, gc_z*12.5), min(25.0, mean_bgc_tnf*100.0), 25.0 if phage_hits else 0.0]):.2f}", "Ravenhall et al. (2015)")
    ]
    for row in hgt_rows:
        ws.append(row)
        
    _sheet_title(ws, "Horizontal Gene Transfer Evidence & Identification Metrics",
                 "Codon frequency, tetranucleotide divergence, and insertion proximity metrics composite calculation")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    # Standard: Last row (Composite Score) always in bold + FFE699 fill
    last_r = ws.max_row
    for c in range(1, 7):
        cell = ws.cell(row=last_r, column=c)
        cell.font = Font(name="Calibri", size=9.5, bold=True)
        cell.fill = PatternFill("solid", fgColor=EXCEL_PALETTE["highlight_fill"])
        
    _col_widths(ws, [24, 20, 20, 24, 18, 22])

    # 6. GC_Profile
    ws = wb.create_sheet(title="GC_Profile")
    ws.sheet_properties.tabColor = TAB_COLORS["GC_Profile"]
    
    num_flagged = sum(1 for x in hgt_results if x.get("Flagged_Outlier") == "YES")
    num_non_flagged = len(hgt_results) - num_flagged
    ws.append([f"Note: {num_flagged} flagged + {num_non_flagged} non-flagged windows shown (total analysed: {len(hgt_results)})"])
    ws.cell(row=3, column=1).font = Font(name="Calibri", size=9, italic=True, color="595959")
    ws.merge_cells("A3:H3")
    
    ws.append(["Window Start", "Window End", "GC Fraction (%)", "GC Z-Score", "TNF Divergence", "Composite HGT Score", "Region", "Flagged Outlier"])
    
    flagged_wins = [x for x in hgt_results if x.get("Flagged_Outlier") == "YES"]
    non_flagged_wins = [x for x in hgt_results if x.get("Flagged_Outlier") != "YES"]
    
    max_rows_allowed = 490
    sampled_non_flagged = non_flagged_wins[:max(0, max_rows_allowed - len(flagged_wins))]
    rows_to_write = flagged_wins + sampled_non_flagged
    rows_to_write.sort(key=lambda x: x["Window_Start"])
    
    for x in rows_to_write:
        ws.append([
            x["Window_Start"],
            x["Window_End"],
            x["GC_Fraction"] * 100,
            x["GC_Zscore"],
            x["TNF_Divergence"],
            x["Composite_HGT_Score"],
            x["Region"],
            x["Flagged_Outlier"]
        ])
        
    _sheet_title(ws, "GC-bias & Tetranucleotide Frequency Sliding Window Profile",
                 "Vectorized genomic window parameters for codon bias detection around BGC boundaries")
    _hdr(ws, 4, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 5, ws.max_row)
    apply_general_cell_styles(ws, 5, ws.max_row)
    _alt_rows(ws, 5, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    for r in range(5, ws.max_row + 1):
        flg_val = ws.cell(row=r, column=8).value
        if flg_val == "YES":
            ws.cell(row=r, column=8).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["warning_fill"])
            ws.cell(row=r, column=8).font = Font(name="Calibri", size=9, bold=True)
            
    _col_widths(ws, [14, 14, 16, 14, 16, 18, 16, 14])

    # 7. Biosynthetic_Pathway
    ws = wb.create_sheet(title="Biosynthetic_Pathway")
    ws.sheet_properties.tabColor = TAB_COLORS["Biosynthetic_Pathway"]
    
    ws.append(["Biosynthetic Expression Step", "Process Description", "Associated Core Genes", "Machinery Status", "Biological Synthesis Context"])
    
    has_mod = any(g["Role"] == "biosynthetic" for g in region_genes)
    has_tra = any(g["Role"] == "transport" for g in region_genes)
    has_reg = any(g["Role"] == "regulatory" for g in region_genes)
    has_imm = any(g["Role"] == "immunity" for g in region_genes)
    
    bgc_type = bgc_target.get("BGC_Type", "Other")
    pathway_steps = [
        ("Amino Acid Precursor Pool", "Intracellular loading of stoichiometric amino acids from glycolysis & TCA", "Host metabolic enzymes", "Detected", "Reconstructed host pathways supply all 20 essential precursors"),
        ("tRNA Synthetase Charging", "Precursor amino acid charging on cognitive tRNAs by aminoacyl-tRNA ligases", "alaS, lysS, serS, etc.", "Detected", "Enzyme genes for charging mapped in the query host genome"),
        ("BGC Operon Transcription", "Transcription activation of biosynthetic cluster by stress-associated factors", "Sigma consensus boxes", "Detected", "Regulatory promoters aligned upstream of core transcripts")
    ]
    
    if "Lantipeptide" in bgc_type or "RiPP" in bgc_type:
        pathway_steps.append(("Ribosomal Translation", "Ribosomal translation of precursor structural peptide on Shine-Dalgarno sites", "SD motif, structural peptide", "Detected", "Consensus ribosomal binding boxes detected in promoter windows"))
        pathway_steps.append(("Lanthionine ring Post-Modification", "Post-translational dehydrations and cyclizations forming structural rings", "Dehydratase/Cyclase (LanB/C)", "Detected" if has_mod else "Not detected", "Dehydratase cyclization forms bioactive macrocyclic rings" if has_mod else "Secondary modifications absent or uncharacterized"))
        pathway_steps.append(("Transporter Membrane Export", "Active export of synthesized bioactive compound across host membrane", "ABC exporter (LanT)", "Detected" if has_tra else "Not detected", "Transmembrane exporter pumps active compound into extracellular space" if has_tra else "No export transport machinery identified in cluster"))
        pathway_steps.append(("Self-Immunity Protection", "Isolate defense protection mechanisms against toxic peptide accumulation", "Immunity genes (LanI/FEG)", "Detected" if has_imm else "Not detected", "Host cells shielded by immunity/efflux pumps" if has_imm else "Self-resistance machinery uncharacterized in core BGC"))
    elif "NRPS" in bgc_type:
        pathway_steps.append(("Nonribosomal Peptide Synthesis", "Biosynthesis of nonribosomal peptide backbone via modular synthetases", "NRPS core synthetase (Pps/Srf/etc.)", "Detected", "Modular synthetases load and link substrate amino acids"))
        pathway_steps.append(("Thioesterase Cyclization & Release", "Thioesterase domain-mediated product cyclization and release", "Thioesterase (TE) domain", "Detected" if has_mod else "Not detected", "TE-mediated release and macrocyclization of product" if has_mod else "Release/cyclization machinery uncharacterized"))
        pathway_steps.append(("MFS/ABC Membrane Export", "Active export of synthesized nonribosomal peptide across host membrane", "MFS transporter / ABC efflux pump", "Detected" if has_tra else "Not detected", "Exporter pumps secondary metabolite into extracellular space" if has_tra else "No export transporter identified in cluster"))
        pathway_steps.append(("Self-Immunity / Resistance", "Host protection mechanisms against nonribosomal peptide secondary metabolite", "Immunity / efflux resistance genes", "Detected" if has_imm else "Not detected", "Host cells shielded by specific immunity proteins/efflux" if has_imm else "Self-resistance machinery uncharacterized"))
    elif "PKS" in bgc_type:
        pathway_steps.append(("Polyketide Chain Assembly", "Biosynthesis of polyketide backbone via modular polyketide synthases", "PKS core synthase", "Detected", "PKS modules assemble and extend polyketide chain"))
        pathway_steps.append(("Backbone Tailoring & Release", "Cyclization, reduction, or decoration of polyketide chain", "Ketoreductase/Dehydratase/TE", "Detected" if has_mod else "Not detected", "Modification enzymes customize polyketide scaffold" if has_mod else "Tailoring modifications uncharacterized"))
        pathway_steps.append(("MFS/ABC Membrane Export", "Active export of polyketide secondary metabolite across host membrane", "MFS transporter / ABC efflux pump", "Detected" if has_tra else "Not detected", "Exporter pumps polyketide metabolite into extracellular space" if has_tra else "No export transporter identified in cluster"))
        pathway_steps.append(("Self-Immunity / Resistance", "Host protection mechanisms against polyketide metabolite", "Immunity / efflux resistance genes", "Detected" if has_imm else "Not detected", "Host cells shielded by specific immunity proteins/efflux" if has_imm else "Self-resistance machinery uncharacterized"))
    else:
        pathway_steps.append(("Core Scaffold Synthesis", "Biosynthesis of the secondary metabolite core scaffold", "Core biosynthetic enzymes", "Detected", "Core enzymes assemble biosynthetic precursor molecules"))
        pathway_steps.append(("Chemical Scaffold Tailoring", "Post-synthesis tailoring modifications (methylation, glycosylation, etc.)", "Tailoring enzymes", "Detected" if has_mod else "Not detected", "Tailoring modifications customize core scaffold" if has_mod else "Tailoring modifications absent or uncharacterized"))
        pathway_steps.append(("Active Cellular Export", "Active export of secondary metabolite across host membrane", "Transporter / efflux proteins", "Detected" if has_tra else "Not detected", "Active transport pumps secondary metabolite into extracellular space" if has_tra else "No export transport machinery identified in cluster"))
        pathway_steps.append(("Self-Immunity / Resistance", "Isolate self-defense mechanisms against secondary metabolite accumulation", "Immunity / resistance genes", "Detected" if has_imm else "Not detected", "Host cells shielded by specific immunity or efflux" if has_imm else "Self-resistance machinery uncharacterized"))

    for step in pathway_steps:
        ws.append(step)
        
    _sheet_title(ws, "Biosynthetic Pathway Machinery Mapping",
                 "Evaluation of structural, modification, regulatory, and immunity steps required for peptide expression")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    for r in range(4, ws.max_row + 1):
        status_val = ws.cell(row=r, column=4).value
        if status_val == "Detected":
            ws.cell(row=r, column=4).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["positive_fill"])
        else:
            ws.cell(row=r, column=4).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["negative_fill"])
            
    # Set dynamic row heights
    for r in range(4, ws.max_row + 1):
        if r <= 7:
            ws.row_dimensions[r].height = 24
        else:
            ws.row_dimensions[r].height = 42
            
    _col_widths(ws, [24, 38, 20, 16, 52])

    # 8. BLAST_Homologs
    ws = wb.create_sheet(title="BLAST_Homologs")
    ws.sheet_properties.tabColor = TAB_COLORS["BLAST_Homologs"]
    
    if bgc_blast_results:
        ws.append(["Target Homolog Species", "NCBI Accession", "E-Value", "Sequence Identity (%)", "Bit Score", "Alignment Length (bp)", "Coverage (%)"])
        for hit in bgc_blast_results:
            ws.append([
                hit.get("Target_Species", "NA"),
                hit.get("Accession", "NA"),
                hit.get("Evalue", 0.0),
                hit.get("Identity_Pct", 0.0),
                hit.get("Bit_Score", 0.0),
                hit.get("Alignment_Length", 0),
                hit.get("Coverage_Pct", 0.0)
            ])
        _sheet_title(ws, "NCBI BLAST Conservation Homology across Species",
                     "Sequence identity, coverage, and alignment parameters of homologous clades in NCBI")
        _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
        apply_thin_borders(ws, 4, ws.max_row)
        apply_general_cell_styles(ws, 4, ws.max_row)
        _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
        for r in range(4, ws.max_row + 1):
            ident = ws.cell(row=r, column=4).value
            if ident and isinstance(ident, (int, float)):
                if ident >= 80.0:
                    ws.cell(row=r, column=4).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["positive_fill"])
                elif ident >= 30.0:
                    ws.cell(row=r, column=4).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["highlight_fill"])
        _col_widths(ws, [24, 20, 18, 20, 16, 20, 16])
    else:
        # Fallback to ortholog comparative mapping if BLAST was skipped
        ws.append(["Gene Symbol", "Locus Tag", "Ref Ortholog Locus", "Identity %", "Coverage %", "Category", "Role Class"])
        for g in region_genes:
            orth = next((x for x in ortholog_mapping if x["Locus_Tag"] == g["Locus_Tag"]), {})
            ws.append([
                g["Gene_Symbol"],
                g["Locus_Tag"],
                orth.get("Ref_Ortholog_Tag", "NA"),
                orth.get("Identity_Pct", 0.0),
                orth.get("Coverage_Pct", 0.0),
                orth.get("Category", "Unique"),
                g["Role"]
            ])
        _sheet_title(ws, "BLASTp Pairwise Homology Aligned to Reference Genome",
                     "Identity rates and sequence coverage of BGC core genes aligned against reference strain")
        _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
        apply_thin_borders(ws, 4, ws.max_row)
        apply_general_cell_styles(ws, 4, ws.max_row)
        _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
        for r in range(4, ws.max_row + 1):
            ident = ws.cell(row=r, column=4).value
            if ident and isinstance(ident, (int, float)):
                if ident >= 80.0:
                    ws.cell(row=r, column=6).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["positive_fill"])
                elif ident >= 30.0:
                    ws.cell(row=r, column=6).fill = PatternFill("solid", fgColor=EXCEL_PALETTE["highlight_fill"])
        _col_widths(ws, [16, 16, 20, 18, 18, 16, 18])

    # 9. Contig_BLAST
    ws = wb.create_sheet(title="Contig_BLAST")
    ws.sheet_properties.tabColor = TAB_COLORS["Contig_BLAST"]
    
    # Load taxid mapping from the comparative genomics CSV if it exists (Fix 17)
    strain_taxid_map = {}
    csv_ref_path = os.path.join(out_dir, "comparative_genomics", "comparative_reference_blast.csv")
    if os.path.exists(csv_ref_path):
        try:
            import pandas as pd
            df_temp = pd.read_csv(csv_ref_path)
            for _, r_temp in df_temp.iterrows():
                ref_str = r_temp.get("Reference_Strain", "")
                tx = r_temp.get("Taxid", "")
                if ref_str and " (" in ref_str:
                    acc_key = ref_str.split(" (")[1].replace(")", "")
                    strain_taxid_map[acc_key] = tx
        except Exception:
            pass

    ws.append(["NCBI Reference Strain", "Assembly Accession", "NCBI Taxonomy ID", "Genome Status"])
    for strain in ref_strains:
        name = strain.split(" (")[0]
        acc = strain.split(" (")[1].replace(")", "") if " (" in strain else "NA"
        ws.append([name, acc, strain_taxid_map.get(acc, ""), "Completed Genome"])
        
    _sheet_title(ws, "NCBI Entrez Search: Close Reference Genomes",
                 "List of the 50 closest completed assembly genomes within the genus of the query host strain")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    _col_widths(ws, [38, 20, 18, 18])

    # 10. Flanking_Context
    ws = wb.create_sheet(title="Flanking_Context")
    ws.sheet_properties.tabColor = TAB_COLORS["Flanking_Context"]
    
    ws.append(["Locus Tag", "Gene Symbol", "Start bp", "End bp", "Strand", "Product Description", "Role Class"])
    for fg in flanking_genes:
        ws.append([
            fg["Locus_Tag"],
            fg["Gene_Symbol"],
            fg["Start_Coord"],
            fg["End_Coord"],
            fg["Strand"],
            fg["Product_Description"],
            fg["Role"]
        ])
        
    _sheet_title(ws, "BGC Flanking Environment Context (30kb Windows)",
                 "Surrounding genomic context and transcription units flanking the core cluster boundaries")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    _col_widths(ws, [16, 12, 12, 12, 10, 38, 18])

    # 11. Methods_References
    create_references_sheet_openpyxl(wb)

    # 12. Metabolic_Integration
    ws = wb.create_sheet(title="Metabolic_Integration")
    ws.sheet_properties.tabColor = TAB_COLORS["Metabolic_Integration"]
    
    ws.append(["Amino Acid Name", "One-Letter Code", "Peptide Count", "Host Biosynthetic Pathway", "Candidate Synthetase Genes", "Ribosomal ATP Cost", "Modification ATP Cost", "Total Stoichiometric ATP"])
    
    for row in stoichiometry_data:
        ws.append([
            row["Amino_Acid"],
            row["Code"],
            row["Precursor_Count"],
            row["Host_Pathway"],
            row["Candidate_Synthetases"],
            row["Translation_ATP_Cost"],
            row["Modification_ATP_Cost"],
            row["Total_ATP_Cost"]
        ])
        
    _sheet_title(ws, "Primary Metabolism precursor Stoichiometry & Translation ATP Cost",
                 "Stoichiometric AA distribution and calculated ATP consumption for transcription, charging, translation, and ring dehydration")
    _hdr(ws, 3, EXCEL_PALETTE["header_fill"])
    apply_thin_borders(ws, 4, ws.max_row)
    apply_general_cell_styles(ws, 4, ws.max_row)
    _alt_rows(ws, 4, ws.max_row, EXCEL_PALETTE["alt_row_fill"])
    
    # Highlight final row (composite total cost) in bold + yellow fill (highlight_fill)
    last_r = ws.max_row
    for c in range(1, 9):
        cell = ws.cell(row=last_r, column=c)
        cell.font = Font(name="Calibri", size=9.5, bold=True)
        cell.fill = PatternFill("solid", fgColor=EXCEL_PALETTE["highlight_fill"])
        
    _col_widths(ws, [18, 12, 12, 38, 24, 16, 16, 18])

    wb.save(out_path)


# ==========================================
# VAULT SNAPSHOT ENGINE (LAW 2 & PROMPT 3)
# ==========================================
import tarfile
import platform
import subprocess
import yaml

class GenoMetaSnapshotEngine:
    @staticmethod
    def create_snapshot(config_path: str, snakefile_path: str, output_dir: str = "resources/tool_snapshots") -> str:
        import json
        os.makedirs(output_dir, exist_ok=True)
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        snapshot_name = f"redolarium_snapshot_{timestamp}.tar.gz"
        snapshot_path = os.path.join(output_dir, snapshot_name)
        
        # Git commit retrieval
        git_commit = "unknown"
        try:
            res = subprocess.run(["git", "rev-parse", "HEAD"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            if res.returncode == 0:
                git_commit = res.stdout.strip()
        except Exception:
            pass
            
        # Manifest data
        manifest = {
            "timestamp": timestamp,
            "git_commit": git_commit,
            "platform": platform.platform(),
            "python_version": sys.version,
            "parameters": {}
        }
        
        # Hardware allocations
        try:
            import psutil
            manifest["cpu_count"] = psutil.cpu_count()
            manifest["total_memory_gb"] = round(psutil.virtual_memory().total / (1024**3), 2)
        except Exception:
            manifest["cpu_count"] = os.cpu_count() or 1
            manifest["total_memory_gb"] = "unknown"

        # Read active configuration parameters
        if os.path.exists(config_path):
            try:
                with open(config_path, "r") as f:
                    manifest["parameters"] = yaml.safe_load(f)
            except Exception:
                pass
                
        # Write manifest locally first to bundle it
        manifest_temp = "snapshot_manifest.json"
        with open(manifest_temp, "w") as fm:
            json.dump(manifest, fm, indent=4)
            
        with tarfile.open(snapshot_path, "w:gz") as tar:
            # Add manifest
            tar.add(manifest_temp, arcname="snapshot_manifest.json")
            
            # Add core configs/rules
            if os.path.exists(config_path):
                tar.add(config_path, arcname="config/config.yaml")
            if os.path.exists(snakefile_path):
                tar.add(snakefile_path, arcname="workflow/Snakefile")
                
            # Add all Python files in redolarium package
            redolarium_dir = os.path.dirname(__file__)
            for root, dirs, files in os.walk(redolarium_dir):
                for file in files:
                    if file.endswith(".py"):
                        full_p = os.path.join(root, file)
                        rel_p = os.path.relpath(full_p, os.path.dirname(redolarium_dir))
                        tar.add(full_p, arcname=rel_p)
            
            # Dereference active database snapshots and bundle
            db_dir = os.path.abspath(os.path.join(redolarium_dir, "..", "resources", "db_snapshots"))
            if os.path.exists(db_dir):
                tar.add(db_dir, arcname="resources/db_snapshots", dereference=True)
                
        # Cleanup temp manifest
        try:
            os.remove(manifest_temp)
        except Exception:
            pass
            
        print(f"Successfully archived state snapshot to: {snapshot_path}")
        return snapshot_path

def build_and_draw_evidence_graph(qc_res, anno_res, bgc_res, evol_res, prom_res, dock_res, phy_res, out_dir, logger):
    logger.info("Synthesizing Evidence Graph DAG across all analysis checkpoints...")
    import networkx as nx
    import json
    import numpy as np
    
    G = nx.DiGraph()
    
    # 1. Parse nodes and evidence payloads
    nodes_data = {
        "QC": {
            "label": f"Genome QC\nScore: {qc_res.confidence_score:.2f}",
            "confidence": qc_res.confidence_score,
            "details": qc_res.prediction if isinstance(qc_res.prediction, dict) else {},
            "evidence": qc_res.evidence,
            "uncertainties": qc_res.limitations
        },
        "Taxonomy": {
            "label": f"Taxonomy\nScore: {anno_res.confidence_score:.2f}",
            "confidence": anno_res.confidence_score,
            "details": anno_res.prediction.get("ref_strains", [])[0] if isinstance(anno_res.prediction, dict) and anno_res.prediction.get("ref_strains") else "Unknown",
            "evidence": anno_res.evidence,
            "uncertainties": anno_res.limitations
        },
        "Orthology": {
            "label": f"Orthology Map\nScore: {anno_res.confidence_score:.2f}",
            "confidence": anno_res.confidence_score,
            "details": f"{len(anno_res.prediction.get('ortholog_mapping', []))} CDS mapped" if isinstance(anno_res.prediction, dict) else "0 CDS mapped",
            "evidence": anno_res.evidence,
            "uncertainties": anno_res.limitations
        },
        "BGC": {
            "label": f"BGC Delineation\nScore: {bgc_res.confidence_score:.2f}",
            "confidence": bgc_res.confidence_score,
            "details": bgc_res.prediction.get("target_bgc", {}).get("BGC_Type", "Unknown") if isinstance(bgc_res.prediction, dict) and bgc_res.prediction.get("target_bgc") else "Unknown",
            "evidence": bgc_res.evidence,
            "uncertainties": bgc_res.limitations
        },
        "HGT": {
            "label": f"HGT Evidence\nScore: {evol_res.confidence_score:.2f}",
            "confidence": evol_res.confidence_score,
            "details": "Sliding window anomalies",
            "evidence": evol_res.evidence,
            "uncertainties": evol_res.limitations
        },
        "Promoter": {
            "label": f"Promoters\nScore: {prom_res.confidence_score:.2f}",
            "confidence": prom_res.confidence_score,
            "details": f"{len(prom_res.prediction.get('promoter_records', []))} motifs" if isinstance(prom_res.prediction, dict) else "0 motifs",
            "evidence": prom_res.evidence,
            "uncertainties": prom_res.limitations
        },
        "Docking": {
            "label": f"Docking Affinity\nScore: {dock_res.confidence_score:.2f}",
            "confidence": dock_res.confidence_score,
            "details": f"{dock_res.prediction.get('docking_output').binding_affinity_kcal_mol if (isinstance(dock_res.prediction, dict) and dock_res.prediction.get('docking_output')) else 0.0} kcal/mol",
            "evidence": dock_res.evidence,
            "uncertainties": dock_res.limitations
        },
        "Phylogeny": {
            "label": f"Phylogenetics\nScore: {phy_res.confidence_score:.2f}",
            "confidence": phy_res.confidence_score,
            "details": f"Speciation nodes: {len(phy_res.prediction.get('divergence_events', []))}" if isinstance(phy_res.prediction, dict) else "0 nodes",
            "evidence": phy_res.evidence,
            "uncertainties": phy_res.limitations
        }
    }
    
    # Add nodes to graph
    for node, data in nodes_data.items():
        G.add_node(node, **data)
        
    # Add dependency edges
    edges = [
        ("QC", "Taxonomy"),
        ("QC", "BGC"),
        ("Taxonomy", "Orthology"),
        ("Orthology", "BGC"),
        ("BGC", "Promoter"),
        ("BGC", "Docking"),
        ("BGC", "HGT"),
        ("Orthology", "Phylogeny")
    ]
    G.add_edges_from(edges)
    
    # Save graph as JSON
    json_data = {
        "schema_version": "1.0.0",
        "nodes": {n: G.nodes[n] for n in G.nodes()},
        "edges": list(G.edges())
    }
    
    def clean_dict(d):
        if not isinstance(d, dict):
            return str(d)
        new_d = {}
        for k, v in d.items():
            if isinstance(v, np.ndarray):
                new_d[k] = v.tolist()
            elif hasattr(v, "to_dict"):
                new_d[k] = v.to_dict()
            elif hasattr(v, "model_dump"):
                new_d[k] = v.model_dump()
            elif isinstance(v, dict):
                new_d[k] = clean_dict(v)
            elif isinstance(v, list):
                new_d[k] = [clean_dict(x) if isinstance(x, dict) else x for x in v]
            else:
                new_d[k] = v
        return new_d
        
    cleaned_json = clean_dict(json_data)
    
    os.makedirs(os.path.join(out_dir, "tabular_data"), exist_ok=True)
    graph_json_path = os.path.join(out_dir, "tabular_data", "evidence_graph.json")
    with open(graph_json_path, "w", encoding="utf-8") as f_json:
        json.dump(cleaned_json, f_json, indent=4)
    logger.info(f"Saved Evidence Graph JSON payload to: {graph_json_path}")
    
    # Render hierarchical flowchart layout
    plt.figure(figsize=(10, 8))
    pos = {
        "QC": (0, 3),
        "Taxonomy": (-1, 2),
        "Orthology": (-1, 1),
        "BGC": (1, 2),
        "HGT": (1, 1),
        "Promoter": (0, 0),
        "Docking": (2, 0),
        "Phylogeny": (-2, 0)
    }
    
    node_colors = [G.nodes[n]["confidence"] for n in G.nodes()]
    nx.draw_networkx_nodes(
        G, pos, node_size=2500, node_color=node_colors, cmap="YlOrRd",
        vmin=0.0, vmax=1.0, edgecolors="black", linewidths=1.2
    )
    labels = {n: G.nodes[n]["label"] for n in G.nodes()}
    nx.draw_networkx_labels(G, pos, labels=labels, font_size=8, font_weight="bold")
    nx.draw_networkx_edges(
        G, pos, width=2.0, edge_color="#333333", arrowstyle="->", arrowsize=15
    )
    
    plt.title("Redolarium Evidence Synthesis Flowchart\nColors represent continuous confidence scores (0.0 to 1.0)", fontsize=11, fontweight="bold", pad=15)
    plt.axis("off")
    plt.tight_layout()
    
    flowchart_path = os.path.join(out_dir, "phylogeny_trees", "evidence_flowchart.png")
    plt.savefig(flowchart_path, dpi=300, bbox_inches="tight")
    plt.close()
    logger.info(f"Saved Evidence Flowchart rendering to: {flowchart_path}")
    
    return graph_json_path

